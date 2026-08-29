#!/usr/bin/env python3
"""Build Complete Book One directly as an editable Microsoft Word DOCX.

The source HTML is produced by ``build-complete-book.mjs`` from the canonical
Markdown. This script maps that structure to native Word paragraphs, styles,
sections, bookmarks, hyperlinks, fields, headers, footers, and an embedded
cover. No office-suite conversion step is involved.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile

try:
    import math2docx
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
    from lxml import etree
except ImportError as error:
    raise SystemExit(
        "The Word build dependencies are required. Install "
        "frontend/scripts/requirements-word.txt "
        "in a virtual environment before running this script."
    ) from error


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_COVER = (
    REPOSITORY_ROOT
    / "docs/blueprint/book-one-complete/v4-working/assets/complete-book-one-v4-cover.png"
)

BOOK_TITLE = "Consciousness: A Digital Organism"
BOOK_SUBTITLE = "Foundations, Agency, and Research"
BOOK_HEADER = "DIGITAL ORGANISM THEORY  ·  COMPLETE BOOK ONE"
TEAL = RGBColor(0x16, 0x70, 0x6F)
DARK_TEAL = RGBColor(0x17, 0x4D, 0x4B)
INK = RGBColor(0x17, 0x25, 0x22)
MUTED = RGBColor(0x5F, 0x6F, 0x6C)
LIGHT_TEAL_HEX = "EEF5F4"


@dataclass
class HtmlNode:
    tag: str
    attrs: dict[str, str] = field(default_factory=dict)
    children: list[HtmlNode | str] = field(default_factory=list)

    @property
    def classes(self) -> set[str]:
        return set(self.attrs.get("class", "").split())

    def text(self) -> str:
        return "".join(
            child if isinstance(child, str) else child.text() for child in self.children
        )


class TreeParser(HTMLParser):
    VOID_TAGS = {"br", "hr", "img", "meta", "link", "input"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = HtmlNode("document")
        self.stack = [self.root]

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        node = HtmlNode(tag.lower(), {name: value or "" for name, value in attrs})
        self.stack[-1].children.append(node)
        if tag.lower() not in self.VOID_TAGS:
            self.stack.append(node)

    def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self.handle_starttag(tag, attrs)
        if tag.lower() not in self.VOID_TAGS:
            self.handle_endtag(tag)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        for index in range(len(self.stack) - 1, 0, -1):
            if self.stack[index].tag == tag:
                del self.stack[index:]
                return

    def handle_data(self, data: str) -> None:
        self.stack[-1].children.append(data)


@dataclass(frozen=True)
class InlineState:
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass
class RenderContext:
    reference_entry: bool = False
    suppress_body_indent: bool = False


def descendants(node: HtmlNode, tag: str) -> Iterable[HtmlNode]:
    for child in node.children:
        if not isinstance(child, HtmlNode):
            continue
        if child.tag == tag:
            yield child
        yield from descendants(child, tag)


def set_cell_font(font, name: str) -> None:
    font.name = name
    if font._element is not None:
        font._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def get_or_add_paragraph_style(document: Document, name: str, base: str = "Normal"):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    return style


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    set_cell_font(normal.font, "Georgia")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.widow_control = True

    body = get_or_add_paragraph_style(document, "DOT Body")
    set_cell_font(body.font, "Georgia")
    body.font.size = Pt(10.5)
    body.font.color.rgb = INK
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Inches(0.18)
    body.paragraph_format.space_after = Pt(2.5)
    body.paragraph_format.line_spacing = 1.2
    body.paragraph_format.widow_control = True

    heading_1 = styles["Heading 1"]
    set_cell_font(heading_1.font, "Arial")
    heading_1.font.size = Pt(25)
    heading_1.font.bold = True
    heading_1.font.color.rgb = INK
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(18)
    heading_1.paragraph_format.line_spacing = 1.04
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.keep_together = True

    heading_2 = styles["Heading 2"]
    set_cell_font(heading_2.font, "Georgia")
    heading_2.font.size = Pt(14)
    heading_2.font.bold = True
    heading_2.font.color.rgb = DARK_TEAL
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True
    heading_2.paragraph_format.keep_together = True

    heading_3 = styles["Heading 3"]
    set_cell_font(heading_3.font, "Arial")
    heading_3.font.size = Pt(10.8)
    heading_3.font.bold = True
    heading_3.font.color.rgb = DARK_TEAL
    heading_3.paragraph_format.space_before = Pt(12)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True

    label = get_or_add_paragraph_style(document, "DOT Chapter Label")
    set_cell_font(label.font, "Arial")
    label.font.size = Pt(8.3)
    label.font.bold = True
    label.font.all_caps = True
    label.font.color.rgb = TEAL
    label.paragraph_format.space_before = Pt(30)
    label.paragraph_format.space_after = Pt(7)
    label.paragraph_format.keep_with_next = True

    eyebrow = get_or_add_paragraph_style(document, "DOT Book Eyebrow")
    set_cell_font(eyebrow.font, "Arial")
    eyebrow.font.size = Pt(8.5)
    eyebrow.font.bold = True
    eyebrow.font.all_caps = True
    eyebrow.font.color.rgb = TEAL
    eyebrow.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_before = Pt(64)
    eyebrow.paragraph_format.space_after = Pt(18)

    title = get_or_add_paragraph_style(document, "DOT Book Title")
    set_cell_font(title.font, "Georgia")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.02
    title.paragraph_format.space_after = Pt(14)

    subtitle = get_or_add_paragraph_style(document, "DOT Book Subtitle")
    set_cell_font(subtitle.font, "Georgia")
    subtitle.font.size = Pt(15)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x3F, 0x55, 0x51)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)

    author = get_or_add_paragraph_style(document, "DOT Book Author")
    set_cell_font(author.font, "Georgia")
    author.font.size = Pt(12.5)
    author.font.color.rgb = INK
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(20)

    edition = get_or_add_paragraph_style(document, "DOT Book Edition")
    set_cell_font(edition.font, "Arial")
    edition.font.size = Pt(8.7)
    edition.font.color.rgb = MUTED
    edition.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.paragraph_format.space_before = Pt(42)
    edition.paragraph_format.line_spacing = 1.4

    contents_title = get_or_add_paragraph_style(document, "DOT Contents Title")
    set_cell_font(contents_title.font, "Arial")
    contents_title.font.size = Pt(22)
    contents_title.font.bold = True
    contents_title.font.color.rgb = INK
    contents_title.paragraph_format.space_after = Pt(18)
    contents_title.paragraph_format.keep_with_next = True

    equation = get_or_add_paragraph_style(document, "DOT Equation")
    set_cell_font(equation.font, "Cambria Math")
    equation.font.size = Pt(11)
    equation.font.color.rgb = INK
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(4)
    equation.paragraph_format.keep_together = True

    equation_number = get_or_add_paragraph_style(document, "DOT Equation Number")
    set_cell_font(equation_number.font, "Arial")
    equation_number.font.size = Pt(8.5)
    equation_number.font.color.rgb = MUTED
    equation_number.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    equation_number.paragraph_format.space_before = Pt(0)
    equation_number.paragraph_format.space_after = Pt(0)

    reference = get_or_add_paragraph_style(document, "DOT Reference")
    set_cell_font(reference.font, "Georgia")
    reference.font.size = Pt(9.5)
    reference.font.color.rgb = INK
    reference.paragraph_format.left_indent = Inches(0.18)
    reference.paragraph_format.first_line_indent = Inches(-0.18)
    reference.paragraph_format.space_after = Pt(10)
    reference.paragraph_format.line_spacing = 1.15

    table_header = get_or_add_paragraph_style(document, "DOT Table Header")
    set_cell_font(table_header.font, "Arial")
    table_header.font.size = Pt(8.6)
    table_header.font.bold = True
    table_header.font.color.rgb = DARK_TEAL
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.05

    table_body = get_or_add_paragraph_style(document, "DOT Table Body")
    set_cell_font(table_body.font, "Georgia")
    table_body.font.size = Pt(8.7)
    table_body.font.color.rgb = INK
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.08

    callout = get_or_add_paragraph_style(document, "DOT Book Callout")
    set_cell_font(callout.font, "Georgia")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor(0x2E, 0x4E, 0x4A)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(0)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.18

    for list_name in ("List Bullet", "List Number"):
        list_style = styles[list_name]
        set_cell_font(list_style.font, "Georgia")
        list_style.font.size = Pt(10.3)
        list_style.font.color.rgb = INK
        list_style.paragraph_format.space_after = Pt(2.5)
        list_style.paragraph_format.line_spacing = 1.16


def set_page_geometry(section, *, cover: bool = False) -> None:
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    if cover:
        section.top_margin = Inches(0.06)
        section.bottom_margin = Inches(0.06)
        section.left_margin = Inches(0.07)
        section.right_margin = Inches(0.07)
    else:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.64)
        section.gutter = Inches(0.06)
        section.header_distance = Inches(0.27)
        section.footer_distance = Inches(0.28)


def clear_story(story) -> None:
    for paragraph in story.paragraphs:
        paragraph.clear()


def set_field(paragraph, instruction: str, result: str = ""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    paragraph._p.append(OxmlElement("w:r"))
    paragraph._p[-1].append(begin)
    paragraph._p.append(OxmlElement("w:r"))
    paragraph._p[-1].append(instruction_node)
    paragraph._p.append(OxmlElement("w:r"))
    paragraph._p[-1].append(separate)
    result_run = paragraph.add_run(result) if result else None
    paragraph._p.append(OxmlElement("w:r"))
    paragraph._p[-1].append(end)
    return result_run


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run("—  ")
    lead.font.name = "Arial"
    lead.font.size = Pt(8)
    lead.font.color.rgb = MUTED
    set_field(paragraph, " PAGE ")
    tail = paragraph.add_run("  —")
    tail.font.name = "Arial"
    tail.font.size = Pt(8)
    tail.font.color.rgb = MUTED


def add_book_header(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(BOOK_HEADER)
    run.font.name = "Arial"
    run.font.size = Pt(7.2)
    run.font.color.rgb = TEAL


def configure_headers_and_footers(document: Document, body_section) -> None:
    settings = document.settings._element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))
    if settings.find(qn("w:autoHyphenation")) is None:
        auto_hyphenation = OxmlElement("w:autoHyphenation")
        auto_hyphenation.set(qn("w:val"), "true")
        settings.append(auto_hyphenation)
    if settings.find(qn("w:consecutiveHyphenLimit")) is None:
        hyphen_limit = OxmlElement("w:consecutiveHyphenLimit")
        hyphen_limit.set(qn("w:val"), "2")
        settings.append(hyphen_limit)
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    body_section.header.is_linked_to_previous = False
    body_section.even_page_header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    body_section.even_page_footer.is_linked_to_previous = False

    clear_story(body_section.header)
    clear_story(body_section.even_page_header)
    clear_story(body_section.footer)
    clear_story(body_section.even_page_footer)

    add_book_header(body_section.even_page_header.paragraphs[0])
    add_page_number(body_section.footer.paragraphs[0])
    add_page_number(body_section.even_page_footer.paragraphs[0])

    sect_pr = body_section._sectPr
    page_number_type = sect_pr.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        sect_pr.append(page_number_type)
    page_number_type.set(qn("w:start"), "1")


def paragraph_rule(paragraph, color: str = "16706F", size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def style_callout(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_TEAL_HEX)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "16706F")
    borders.append(left)
    p_pr.append(borders)


def bookmark_name(raw: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9_]", "_", raw)
    if not normalized or not normalized[0].isalpha():
        normalized = f"b_{normalized}"
    return normalized[:38]


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def run_properties(run_node, state: InlineState, *, hyperlink: bool = False, superscript=False):
    r_pr = OxmlElement("w:rPr")
    if state.bold:
        r_pr.append(OxmlElement("w:b"))
    if state.italic:
        r_pr.append(OxmlElement("w:i"))
    if state.code:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Consolas")
        fonts.set(qn("w:hAnsi"), "Consolas")
        r_pr.append(fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "19")
        r_pr.append(size)
    if hyperlink:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "16706F")
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
    if superscript:
        vertical = OxmlElement("w:vertAlign")
        vertical.set(qn("w:val"), "superscript")
        r_pr.append(vertical)
    run_node.append(r_pr)


def add_hyperlink(paragraph, text: str, href: str, state: InlineState) -> None:
    if href.startswith("#equation-"):
        match = re.fullmatch(r"(?P<prefix>.*?)(?P<number>\([0-9]+\.[0-9]+\))", text)
        if not match:
            raise RuntimeError(f"Equation cross-reference has an invalid label: {text!r}")
        if match.group("prefix"):
            prefix = paragraph.add_run(match.group("prefix"))
            prefix.bold = state.bold
            prefix.italic = state.italic
        result_run = set_field(
            paragraph,
            f' REF {bookmark_name(href[1:])} \\h ',
            match.group("number"),
        )
        result_run.bold = state.bold
        result_run.italic = state.italic
        result_run.font.color.rgb = TEAL
        result_run.font.underline = True
        return

    hyperlink = OxmlElement("w:hyperlink")
    if href.startswith("#"):
        hyperlink.set(qn("w:anchor"), bookmark_name(href[1:]))
    else:
        relationship_id = paragraph.part.relate_to(
            href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties(
        run,
        state,
        hyperlink=True,
        superscript=href.startswith("#reference-"),
    )
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def normalized_text(value: str) -> str:
    return re.sub(r"\s+", " ", value)


def render_inline(paragraph, children: Iterable[HtmlNode | str], state=InlineState()) -> None:
    for child in children:
        if isinstance(child, str):
            text = normalized_text(child)
            if not text:
                continue
            run = paragraph.add_run(text)
            run.bold = state.bold
            run.italic = state.italic
            if state.code:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK_TEAL
            continue

        if child.tag == "strong":
            render_inline(
                paragraph,
                child.children,
                InlineState(bold=True, italic=state.italic, code=state.code),
            )
        elif child.tag == "em":
            render_inline(
                paragraph,
                child.children,
                InlineState(bold=state.bold, italic=True, code=state.code),
            )
        elif child.tag == "code":
            render_inline(
                paragraph,
                child.children,
                InlineState(bold=state.bold, italic=state.italic, code=True),
            )
        elif child.tag == "span" and "math-inline" in child.classes:
            latex = child.attrs.get("data-latex", "").strip()
            if not latex:
                raise RuntimeError("Inline equation is missing its LaTeX source")
            math2docx.add_math(paragraph, latex)
        elif child.tag == "a":
            add_hyperlink(
                paragraph,
                normalized_text(child.text()),
                child.attrs.get("href", ""),
                state,
            )
        elif child.tag == "br":
            paragraph.add_run().add_break()
        else:
            render_inline(paragraph, child.children, state)


def direct_list_content(node: HtmlNode) -> list[HtmlNode | str]:
    content: list[HtmlNode | str] = []
    for child in node.children:
        if isinstance(child, HtmlNode) and child.tag in {"ul", "ol"}:
            continue
        if isinstance(child, HtmlNode) and child.tag == "p":
            content.extend(child.children)
        else:
            content.append(child)
    return content


def render_list(
    document: Document,
    node: HtmlNode,
    context: RenderContext,
    *,
    level: int = 0,
) -> None:
    ordered = node.tag == "ol"
    list_style = "List Number" if ordered else "List Bullet"
    for item in (
        child for child in node.children if isinstance(child, HtmlNode) and child.tag == "li"
    ):
        paragraph = document.add_paragraph(style=list_style)
        paragraph.paragraph_format.left_indent = Inches(0.22 + level * 0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(4)
        render_inline(paragraph, direct_list_content(item))
        for nested in (
            child
            for child in item.children
            if isinstance(child, HtmlNode) and child.tag in {"ul", "ol"}
        ):
            render_list(document, nested, context, level=level + 1)
    context.suppress_body_indent = True


def set_cell_margins(cell, *, top=55, start=70, bottom=55, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{edge}"))
        if margin is None:
            margin = OxmlElement(f"w:{edge}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.first_child_found_in("w:tcW")
    tc_width.set(qn("w:w"), str(round(width_inches * 1440)))
    tc_width.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def configure_table_xml(table, *, borders: bool) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "single" if borders else "nil")
        if borders:
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "A8C7C2")


def prevent_row_split(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        tr_pr.append(OxmlElement("w:tblHeader"))


def render_table(document: Document, node: HtmlNode, context: RenderContext) -> None:
    source_rows = list(descendants(node, "tr"))
    row_cells = [
        [
            child
            for child in row.children
            if isinstance(child, HtmlNode) and child.tag in {"th", "td"}
        ]
        for row in source_rows
    ]
    column_count = max((len(cells) for cells in row_cells), default=0)
    if not column_count:
        return

    table = document.add_table(rows=len(row_cells), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    configure_table_xml(table, borders=True)
    if column_count == 3:
        widths = (0.65, 1.95, 1.85)
    else:
        widths = tuple(4.45 / column_count for _ in range(column_count))

    for row_index, (target_row, source_cells) in enumerate(zip(table.rows, row_cells)):
        is_header = bool(source_cells) and all(cell.tag == "th" for cell in source_cells)
        prevent_row_split(target_row, repeat_header=is_header)
        for column_index, target_cell in enumerate(target_row.cells):
            set_cell_width(target_cell, widths[column_index])
            set_cell_margins(target_cell)
            target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_header:
                set_cell_shading(target_cell, "E8F3F1")
            paragraph = target_cell.paragraphs[0]
            paragraph.style = "DOT Table Header" if is_header else "DOT Table Body"
            if column_index < len(source_cells):
                render_inline(paragraph, source_cells[column_index].children)

    context.suppress_body_indent = True


def render_display_equation(
    document: Document,
    node: HtmlNode,
    bookmark_counter: list[int],
    context: RenderContext,
) -> None:
    latex = node.attrs.get("data-latex", "").strip()
    equation_number = node.attrs.get("data-equation-number", "").strip()
    equation_id = node.attrs.get("id", "").strip()
    match = re.fullmatch(r"([0-9]+)\.([0-9]+)", equation_number)
    if not latex or not equation_id or not match:
        raise RuntimeError(f"Malformed display equation metadata: {equation_number!r}")
    chapter_number, local_number = match.groups()

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    configure_table_xml(table, borders=False)
    prevent_row_split(table.rows[0])
    widths = (0.45, 3.55, 0.45)
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        set_cell_margins(cell, top=20, start=15, bottom=20, end=15)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)

    equation_paragraph = table.cell(0, 1).paragraphs[0]
    equation_paragraph.style = "DOT Equation"
    math2docx.add_math(equation_paragraph, latex)

    number_paragraph = table.cell(0, 2).paragraphs[0]
    number_paragraph.style = "DOT Equation Number"
    number_paragraph.add_run(f"({chapter_number}.")
    sequence_run = set_field(
        number_paragraph,
        f" SEQ DOTEquation{chapter_number} \\* ARABIC ",
        local_number,
    )
    sequence_run.font.name = "Arial"
    sequence_run.font.size = Pt(8.5)
    sequence_run.font.color.rgb = MUTED
    number_paragraph.add_run(")")
    add_bookmark(
        number_paragraph,
        bookmark_name(equation_id),
        bookmark_counter[0],
    )
    bookmark_counter[0] += 1
    context.suppress_body_indent = True


def render_block(
    document: Document,
    node: HtmlNode,
    bookmark_counter: list[int],
    context: RenderContext,
) -> None:
    if node.tag in {"h1", "h2", "h3"}:
        level = int(node.tag[-1])
        paragraph = document.add_paragraph(style=f"Heading {level}")
        render_inline(paragraph, node.children)
        if node.attrs.get("id"):
            add_bookmark(paragraph, bookmark_name(node.attrs["id"]), bookmark_counter[0])
            bookmark_counter[0] += 1
        if level == 1:
            paragraph_rule(paragraph, color="A8C7C2", size="6")
        context.reference_entry = level == 3 and bool(
            re.fullmatch(r"Reference [0-9]+", node.text().strip())
        )
        context.suppress_body_indent = True
        return

    if node.tag == "p":
        is_label = "chapter-label" in node.classes
        if is_label:
            style = "DOT Chapter Label"
        elif context.reference_entry:
            style = "DOT Reference"
        else:
            style = "DOT Body"
        if not node.text().strip() and "chapter-label" not in node.classes:
            return
        paragraph = document.add_paragraph(style=style)
        if style == "DOT Body" and context.suppress_body_indent:
            paragraph.paragraph_format.first_line_indent = Inches(0)
        render_inline(paragraph, node.children)
        context.suppress_body_indent = is_label
        return

    if node.tag in {"ul", "ol"}:
        render_list(document, node, context)
        return

    if node.tag == "blockquote":
        callout_nodes = [
            child
            for child in node.children
            if isinstance(child, HtmlNode) and child.tag in {"p", "ul", "ol"}
        ]
        for index, child in enumerate(callout_nodes):
            if child.tag in {"ul", "ol"}:
                render_list(document, child, context)
                continue
            paragraph = document.add_paragraph(style="DOT Book Callout")
            render_inline(paragraph, child.children)
            style_callout(paragraph)
            if index == 0:
                paragraph.paragraph_format.space_before = Pt(7)
            if index == len(callout_nodes) - 1:
                paragraph.paragraph_format.space_after = Pt(8)
        context.suppress_body_indent = True
        return

    if node.tag == "div" and "math-display" in node.classes:
        render_display_equation(document, node, bookmark_counter, context)
        return

    if node.tag == "table":
        render_table(document, node, context)
        return

    if node.tag == "hr":
        paragraph = document.add_paragraph()
        paragraph_rule(paragraph)
        context.suppress_body_indent = True
        return

    for child in node.children:
        if isinstance(child, HtmlNode):
            render_block(document, child, bookmark_counter, context)


def add_cover(document: Document, cover_path: Path) -> None:
    section = document.sections[0]
    set_page_geometry(section, cover=True)
    for story in (
        section.header,
        section.even_page_header,
        section.footer,
        section.even_page_footer,
    ):
        story.is_linked_to_previous = False
        clear_story(story)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run.add_picture(str(cover_path), width=Inches(5.86), height=Inches(8.79))
    drawing_properties = run._r.xpath(".//wp:docPr")
    if drawing_properties:
        drawing_properties[0].set(
            "descr", "Cover of Consciousness: A Digital Organism by Henok Ghebrechristos"
        )


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph("Digital Organism Theory · Complete Book One")
    paragraph.style = "DOT Book Eyebrow"

    title = document.add_paragraph(style="DOT Book Title")
    title.add_run("Consciousness:")
    title.add_run().add_break()
    title.add_run("A Digital Organism")

    document.add_paragraph(BOOK_SUBTITLE, style="DOT Book Subtitle")
    rule = document.add_paragraph()
    rule.paragraph_format.left_indent = Inches(1.55)
    rule.paragraph_format.right_indent = Inches(1.55)
    paragraph_rule(rule, size="12")
    document.add_paragraph("Henok Ghebrechristos", style="DOT Book Author")

    edition = document.add_paragraph(style="DOT Book Edition")
    edition.add_run("Editorial Manuscript v4.2 · August 2026")
    edition.add_run().add_break()
    edition.add_run("Prepared for author line edit and manual refinement")
    edition.add_run().add_break()
    edition.add_run("Expanded directly from the immutable Public Reader’s Edition v3")


def add_rights_page(document: Document) -> None:
    first = document.add_paragraph("© 2026 Henok Ghebrechristos. All rights reserved.")
    first.style = "DOT Body"
    first.paragraph_format.space_before = Inches(2.5)
    document.add_paragraph(
        "This is an unreleased editorial manuscript. It is not the public Reader’s "
        "Edition and should not be treated as a sealed publication version.",
        style="DOT Body",
    )
    document.add_paragraph(
        "Prepared after developmental editing for the author’s line edit and final "
        "manual refinement. Observation, established external model, DOT derivation, "
        "and speculation remain explicitly distinguished.",
        style="DOT Body",
    )


def add_contents(document: Document) -> None:
    document.add_paragraph("EDITORIAL MANUSCRIPT", style="DOT Chapter Label")
    document.add_paragraph("Contents", style="DOT Contents Title")
    toc = document.add_paragraph(style="DOT Body")
    toc.paragraph_format.space_before = Pt(6)
    set_field(
        toc,
        ' TOC \\o "1-1" \\h \\z \\u ',
        "Microsoft Word will update this table with chapter titles and page numbers.",
    )


def validate_word_package(output_path: Path, html_root: HtmlNode) -> dict[str, int]:
    display_equations = [
        node
        for node in descendants(html_root, "div")
        if "math-display" in node.classes
    ]
    inline_equations = [
        node
        for node in descendants(html_root, "span")
        if "math-inline" in node.classes
    ]
    links = list(descendants(html_root, "a"))
    equation_crossrefs = [
        link
        for link in links
        if re.fullmatch(r"#equation-[0-9]+-[0-9]+", link.attrs.get("href", ""))
    ]
    external_targets = {
        link.attrs.get("href", "")
        for link in links
        if link.attrs.get("href", "").startswith(("https://", "http://"))
    }

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    with ZipFile(output_path) as package:
        document_xml = etree.fromstring(package.read("word/document.xml"))
        relationships_xml = etree.fromstring(
            package.read("word/_rels/document.xml.rels")
        )
        bookmark_names = document_xml.xpath(
            "//w:bookmarkStart/@w:name", namespaces=namespaces
        )
        duplicate_bookmarks = sorted(
            {name for name in bookmark_names if bookmark_names.count(name) > 1}
        )
        if duplicate_bookmarks:
            raise RuntimeError(
                "Duplicate Word bookmark names: " + ", ".join(duplicate_bookmarks)
            )

        internal_anchors = document_xml.xpath(
            "//w:hyperlink/@w:anchor", namespaces=namespaces
        )
        missing_anchors = sorted(set(internal_anchors) - set(bookmark_names))
        if missing_anchors:
            raise RuntimeError(
                "Internal links target missing bookmarks: " + ", ".join(missing_anchors)
            )

        expected_equation_bookmarks = {
            bookmark_name(node.attrs["id"]) for node in display_equations
        }
        missing_equation_bookmarks = sorted(
            expected_equation_bookmarks - set(bookmark_names)
        )
        if missing_equation_bookmarks:
            raise RuntimeError(
                "Equations are missing Word bookmarks: "
                + ", ".join(missing_equation_bookmarks)
            )

        field_codes = [
            "".join(node.itertext()).strip()
            for node in document_xml.xpath("//w:instrText", namespaces=namespaces)
        ]
        sequence_count = sum(
            code.startswith("SEQ DOTEquation") for code in field_codes
        )
        equation_ref_count = sum(
            code.startswith("REF equation_") for code in field_codes
        )
        if sequence_count != len(display_equations):
            raise RuntimeError(
                f"Expected {len(display_equations)} equation-number fields, "
                f"found {sequence_count}"
            )
        if equation_ref_count != len(equation_crossrefs):
            raise RuntimeError(
                f"Expected {len(equation_crossrefs)} equation-reference fields, "
                f"found {equation_ref_count}"
            )

        native_math_count = len(
            document_xml.xpath("//m:oMath", namespaces=namespaces)
        )
        expected_math_count = len(display_equations) + len(inline_equations)
        if native_math_count != expected_math_count:
            raise RuntimeError(
                f"Expected {expected_math_count} native Word equations, "
                f"found {native_math_count}"
            )

        relationship_namespace = relationships_xml.nsmap.get(None)
        actual_external_targets = {
            relationship.get("Target")
            for relationship in relationships_xml.findall(
                f"{{{relationship_namespace}}}Relationship"
            )
            if relationship.get("TargetMode") == "External"
        }
        if actual_external_targets != external_targets:
            raise RuntimeError("External Word hyperlink relationships do not match source")
        if any(target.startswith("file:") for target in actual_external_targets):
            raise RuntimeError("Word package contains an external local-file relationship")

        package_text = "".join(document_xml.itertext())
        leaked_markers = [
            marker
            for marker in ("DOT_INLINE_MATH", "DOT_DISPLAY_MATH")
            if marker in package_text
        ]
        if leaked_markers:
            raise RuntimeError("Math transport markers leaked into the Word document")
        media_parts = [
            name for name in package.namelist() if name.startswith("word/media/")
        ]
        if not media_parts:
            raise RuntimeError("Word package has no embedded cover image")

    return {
        "display_equations": len(display_equations),
        "inline_equations": len(inline_equations),
        "equation_crossrefs": len(equation_crossrefs),
        "external_links": len(external_targets),
    }


def build_word_document(
    html_path: Path, output_path: Path, cover_path: Path
) -> tuple[int, dict[str, int]]:
    parser = TreeParser()
    parser.feed(html_path.read_text(encoding="utf-8"))
    sections = [section for section in descendants(parser.root, "section")]
    chapter_sections = [section for section in sections if "chapter" in section.classes]
    if len(chapter_sections) != 12:
        raise RuntimeError(f"Expected 12 manuscript sections, found {len(chapter_sections)}")

    document = Document()
    configure_styles(document)
    document.core_properties.title = f"{BOOK_TITLE} — Complete Book One"
    document.core_properties.author = "Henok Ghebrechristos"
    document.core_properties.subject = "Digital Organism Theory"
    document.core_properties.comments = "Editable Microsoft Word editorial manuscript v4.2"
    document.core_properties.keywords = "consciousness, Digital Organism Theory, DOT"

    add_cover(document, cover_path)

    front_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_geometry(front_section)
    for story in (
        front_section.header,
        front_section.even_page_header,
        front_section.footer,
        front_section.even_page_footer,
    ):
        story.is_linked_to_previous = False
        clear_story(story)

    add_title_page(document)
    document.add_page_break()
    add_rights_page(document)
    document.add_page_break()
    add_contents(document)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_geometry(body_section)
    configure_headers_and_footers(document, body_section)

    bookmark_counter = [1]
    for section_index, chapter in enumerate(chapter_sections):
        if section_index:
            document.add_page_break()
        context = RenderContext()
        for child in chapter.children:
            if isinstance(child, HtmlNode):
                render_block(document, child, bookmark_counter, context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    # Reopening exercises the same package relationships Word will consume.
    reopened = Document(output_path)
    if not reopened.inline_shapes:
        raise RuntimeError("Word package was saved without its embedded cover image")
    audit = validate_word_package(output_path, parser.root)
    return len(chapter_sections), audit


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path, help="HTML manuscript from build-complete-book.mjs")
    parser.add_argument("output", type=Path, help="Editable Microsoft Word DOCX")
    parser.add_argument("--cover", type=Path, default=DEFAULT_COVER)
    args = parser.parse_args()

    input_path = args.input.resolve()
    output_path = args.output.resolve()
    cover_path = args.cover.resolve()
    if not input_path.is_file():
        raise SystemExit(f"Input HTML does not exist: {input_path}")
    if not cover_path.is_file():
        raise SystemExit(f"Cover image does not exist: {cover_path}")

    chapter_count, audit = build_word_document(input_path, output_path, cover_path)
    print(
        f"Built {output_path.name}: native Word DOCX, {chapter_count} manuscript "
        f"sections, {audit['display_equations']} numbered display equations, "
        f"{audit['inline_equations']} inline equations, "
        f"{audit['equation_crossrefs']} live equation cross-references, "
        f"{audit['external_links']} external source links"
    )


if __name__ == "__main__":
    main()
